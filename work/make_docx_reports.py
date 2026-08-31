# -*- coding: utf-8 -*-
"""يولّد ملفي وورد (docx) يسردان كل تصحيح أُجري على تفريغ الكتابين،
مرتبًا على صفحات الكتاب، بحيث يسهل الرجوع إلى الصفحة المطبوعة نفسها.

المدخلات:
  partN_decoded_v2.txt  = النص «قبل» التصحيح (التفريغ القديم الخاطئ)
  partN_final.txt       = النص «بعد» التصحيح (المطابق لما هو مطبوع في الكتاب)

الخرج:
  سجل_التصحيحات_الجزء_الثاني.docx
  سجل_التصحيحات_الجزء_الثالث.docx

ترقيم الصفحات: أول صفحة في ملف PDF غلاف/ترقيم «0»، لذا:
  الصفحة المطبوعة في الكتاب = رقم صفحة PDF − 1
(يُتحقَق من هذه العلاقة آليًّا من رقم الصفحة المطبوع داخل كل صفحة.)
"""
import difflib
import re
from collections import Counter, OrderedDict

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "Sakkal Majalla"
MARK = re.compile(r"^===== PAGE (\d+) =====$")
ARABIC = re.compile(r"[\u0621-\u064A]")
FAKE_MARKS = "\u06e1\u06d6"  # ۡۖ

# ---------------------------------------------------------------- تصنيف التغيير


def classify(before: str, after: str) -> str:
    """يحدد نوع التصحيح في السطر بمقارنة قبل/بعد."""
    if before.strip().isdigit() and after.strip().isdigit():
        return "رقم صفحة مطبوع معكوس"
    types = []
    bd, ad = re.findall(r"\d+", before), re.findall(r"\d+", after)
    if bd != ad and any(x and x[::-1] in ad for x in bd):
        types.append("أرقام معكوسة")
    if sorted(before) == sorted(after):
        types.append("موضع رقم/علامة صُحِّح داخل السطر")
    if before.count(FAKE_MARKS) > after.count(FAKE_MARKS):
        types.append("علامات ضبط زائفة حُذفت (رسم المسافة)")
    if len(ARABIC.findall(after)) > len(ARABIC.findall(before)):
        types.append("حروف ساقطة استُعيدت")
    added_punct = [p for p in ":،؛." if after.count(p) > before.count(p)]
    if added_punct:
        types.append("علامة ترقيم ساقطة استُعيدت")
    if not types:
        types.append("ضبط الحركات/الترتيب")
    return " + ".join(dict.fromkeys(types))


# ------------------------------------------------------- قراءة الفروق والصفحات


def page_of(lines, idx):
    """رقم صفحة PDF الذي يقع فيه السطر idx (آخر علامة صفحة قبله)."""
    page = None
    for i in range(idx, -1, -1):
        m = MARK.match(lines[i])
        if m:
            page = int(m.group(1))
            break
    return page


def printed_page_map(final_lines):
    """خريطة رقم صفحة PDF ← الرقم المطبوع المستخرَج من نفس الصفحة."""
    pages, cur = {}, None
    for line in final_lines:
        m = MARK.match(line)
        if m:
            cur = int(m.group(1))
            pages[cur] = None
            continue
        if cur is not None and re.fullmatch(r"\d{1,3}", line.strip()):
            pages[cur] = line.strip()
    return pages


def collect(part: str):
    before = open(f"{part}_decoded_v2.txt", encoding="utf-8").read().splitlines()
    after = open(f"{part}_final.txt", encoding="utf-8").read().splitlines()
    sm = difflib.SequenceMatcher(None, before, after, autojunk=False)
    rows = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue
        for k in range(max(i2 - i1, j2 - j1)):
            bi, ji = i1 + k, j1 + k
            b = before[bi] if bi < i2 else ""
            a = after[ji] if ji < j2 else ""
            if b == a:
                continue
            page = page_of(after, min(ji, len(after) - 1))
            rows.append({"page": page, "before": b, "after": a,
                         "type": classify(b, a)})
    rows.sort(key=lambda r: (r["page"] or 0,))
    return rows


# ---------------------------------------------------------------- أدوات docx


def rtl_par(p, align=WD_ALIGN_PARAGRAPH.RIGHT):
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    p.alignment = align


def style_run(r, size=11, bold=False, color=None):
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    rPr = r._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:cs"), FONT)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(int(size * 2)))
    rPr.append(szCs)
    if bold:
        bCs = OxmlElement("w:bCs")
        bCs.set(qn("w:val"), "1")
        rPr.append(bCs)
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def add_par(doc, text, size=11, bold=False, color=None,
            align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6):
    p = doc.add_paragraph()
    rtl_par(p, align)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color)
    return p


def table_rtl(t):
    tblPr = t._tbl.tblPr
    bv = OxmlElement("w:bidiVisual")
    tblPr.append(bv)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def cell_text(cell, text, size=10, bold=False, color=None,
              align=WD_ALIGN_PARAGRAPH.RIGHT):
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    rtl_par(p, align)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color)


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


# ---------------------------------------------------------------- بناء الملف


def build(part: str, title: str, pdf_name: str, n_pages: int, out: str):
    rows = collect(part)
    final = open(f"{part}_final.txt", encoding="utf-8").read().splitlines()
    pmap = printed_page_map(final)

    # تحقّق آلي من علاقة الترقيم: المطبوعة = PDF − 1
    mism = [pg for pg, pr in pmap.items()
            if pr is not None and int(pr) != pg - 1]
    assert not mism, f"صفحات تخالف القاعدة: {mism}"
    printed = lambda pg: (pg or 1) - 1

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    sec.left_margin = sec.right_margin = Cm(1.4)
    sec.top_margin = sec.bottom_margin = Cm(1.3)

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)

    add_par(doc, f"سجلّ التصحيحات مرتّبًا على صفحات الكتاب — {title}",
            size=19, bold=True, color=(0x1F, 0x4E, 0x79), space_after=4)
    add_par(doc, f"المصدر: «{pdf_name}» ({n_pages} صفحة في ملف PDF — من الغلاف ص 0 إلى ص {n_pages - 2})",
            size=11, color=(0x40, 0x40, 0x40), space_after=10)

    # ------------------------------------------------------------- تمهيد
    intro = [
        "كيف تقرأ هذا الملف؟",
        "١) كل سطر في الجدول الرئيس تصحيحٌ واحد: عمود «قبل» هو ما كان يُخرجه التفريغ القديم الخاطئ (محفوظ في work/" + part + "_decoded_v2.txt)، وعمود «بعد» هو الصواب المطابق لما هو مطبوع في كتابك نفسه (work/" + part + "_final.txt).",
        "٢) المرجع في كل تصحيح هو الكتاب نفسه (ملف PDF الذي فُرِّغ) — أي أن التصحيح أعاد النص إلى المطبوع في الصفحة المذكورة أمامك؛ ولم تُقارن النسخة بأي طبعة خارجية أخرى.",
        "٣) عمود «الصفحة» يعطيك رقم الصفحة المطبوعة في الكتاب، وبين قوسين رقمها في ملف PDF. الصفحة المطبوعة = صفحة PDF − 1، لأن أول صفحة في الملف غلافٌ مرقّم (0). للتحقق من أي سطر: افتح الكتاب على الصفحة المطبوعة وقارن ما فيها بعمود «بعد».",
        "٤) جدول الملخص يبيّن حجم كل نوع من الأخطاء، ثم يليه فهرس الصفحات المتأثرة بعدد التصحيحات في كل صفحة.",
    ]
    for i, t in enumerate(intro):
        add_par(doc, t, size=12 if i == 0 else 11, bold=(i == 0), space_after=4)

    # ------------------------------------------------------------- الملخص
    add_par(doc, "الملخّص الرقمي", size=14, bold=True,
            color=(0x1F, 0x4E, 0x79), space_after=4)
    add_par(doc, f"إجمالي الأسطر المصحّحة: {len(rows)} سطرًا، موزّعة على "
                 f"{len({r['page'] for r in rows})} صفحة من صفحات الكتاب.",
            size=11, space_after=6)

    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    table_rtl(t)
    hdr = t.rows[0]
    repeat_header(hdr)
    cell_text(hdr.cells[0], "نوع التصحيح", bold=True, color=(0xFF, 0xFF, 0xFF))
    cell_text(hdr.cells[1], "عدد الأسطر", bold=True, color=(0xFF, 0xFF, 0xFF),
              align=WD_ALIGN_PARAGRAPH.CENTER)
    for c in hdr.cells:
        shade(c, "1F4E79")
    counts = Counter()
    for r in rows:
        for typ in r["type"].split(" + "):
            counts[typ] += 1
    for typ, n in counts.most_common():
        row = t.add_row()
        cell_text(row.cells[0], typ)
        cell_text(row.cells[1], str(n), align=WD_ALIGN_PARAGRAPH.CENTER)
    t.columns[0].width = Cm(18.0)
    t.columns[1].width = Cm(7.0)
    for row in t.rows:
        row.cells[0].width = Cm(18.0)
        row.cells[1].width = Cm(7.0)

    # --------------------------------------------------- فهرس الصفحات المتأثرة
    by_page = Counter(r["page"] for r in rows)
    add_par(doc, "فهرس الصفحات المتأثرة (الصفحة المطبوعة: عدد التصحيحات)",
            size=14, bold=True, color=(0x1F, 0x4E, 0x79), space_after=4)
    items = "، ".join(f"ص {printed(pg)}: {n}" for pg, n in sorted(by_page.items()))
    add_par(doc, items + ".", size=10, space_after=10)

    # ------------------------------------------------------------- الجدول الرئيس
    add_par(doc, "الجدول الرئيس: التصحيحات سطرًا سطرًا",
            size=14, bold=True, color=(0x1F, 0x4E, 0x79), space_after=4)

    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    table_rtl(t)
    hdr = t.rows[0]
    repeat_header(hdr)
    heads = ["م", "الصفحة (PDF)", "قبل التصحيح (التفريغ الخاطئ)",
             "بعد التصحيح (المطابق للكتاب)", "نوع التصحيح"]
    widths = [Cm(1.1), Cm(2.6), Cm(10.6), Cm(10.6), Cm(2.9)]
    for c, h, w in zip(hdr.cells, heads, widths):
        cell_text(c, h, bold=True, color=(0xFF, 0xFF, 0xFF),
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        shade(c, "1F4E79")
    for n, r in enumerate(rows, 1):
        row = t.add_row()
        cell_text(row.cells[0], str(n), size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[1],
                  f"ص {printed(r['page'])} (PDF {r['page']})", size=9,
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[2], r["before"] or "—(سطر محذوف)—", size=9.5,
                  color=(0x8B, 0x00, 0x00))
        cell_text(row.cells[3], r["after"] or "—(حُذف السطر)—", size=9.5,
                  color=(0x00, 0x60, 0x33))
        cell_text(row.cells[4], r["type"], size=9)
        if r["before"] == "" or r["after"] == "":
            shade(row.cells[2], "FFF2F2") if r["before"] == "" else None
    for row in t.rows:
        for c, w in zip(row.cells, widths):
            c.width = w

    # ------------------------------------------------------------- خاتمة
    doc.add_page_break()
    add_par(doc, "ما لم يُصحَّح ويحتاج مراجعتك البصرية", size=14, bold=True,
            color=(0x1F, 0x4E, 0x79), space_after=4)
    notes = [
        "١) بعض المسافات بين الكلمات لا تزال غير ظاهرة (مثل «قليلاًمع») لأن برنامج الإخراج يستعيض عن المسافة بمحاذاة الحروف؛ لم تُضَف مسافات تخمينية تجنّبًا لتمزيق الكلمات الممدودة.",
        "٢) في الكلمات ذات الرسوم المركّبة قد تقع الحركة على الحرف المجاور (مثل «إِلي» بدل «إِلىٰ») — راجع هذه الصفحات بصريًا عند النشر.",
        "٣) ترتيب خانات الجداول والرسوم وتدوين أرقام الهوامش المتفرّقة يحتاج مطابقة بصرية مع ملف PDF عند الإخراج النهائي.",
        "٤) تصحيحات هذا الملف كلها مقابل الكتاب نفسه (ملف PDF ver3): «قبل» كان قراءة آلية خاطئة للصفحة، و«بعد» هو المطبوع فيها فعلًا، والصفحة مذكورة في كل سطر.",
    ]
    for note in notes:
        add_par(doc, note, size=11, space_after=4)

    doc.save(out)
    return len(rows), counts


if __name__ == "__main__":
    jobs = OrderedDict([
        ("part2", ("الجزء الثاني", "كتاب رتل مصمم ver3 - Part2.pdf", 63,
                   "سجل_التصحيحات_الجزء_الثاني.docx")),
        ("part3", ("الجزء الثالث", "كتاب رتل مصمم ver3 - Part3.pdf", 32,
                   "سجل_التصحيحات_الجزء_الثالث.docx")),
    ])
    for part, (title, pdf, n, out) in jobs.items():
        n_rows, counts = build(part, title, pdf, n, out)
        print(f"{out}: {n_rows} تصحيحًا")
        for typ, c in counts.most_common():
            print(f"   {c:4d}  {typ}")
